# src/parsers/docx_parser.py

"""
DOCX text extraction using python-docx.

Extracts text from:
    - Paragraphs
    - Tables
    - Headers and footers

Why python-docx instead of docx2pdf conversion:
    Direct parsing is faster and does not require Microsoft Word installed.
    docx2pdf conversion was found unreliable across different OS environments.
"""

from pathlib import Path

from src.utils.logger import get_logger

logger = get_logger(__name__)


# Paragraph extraction

def _extract_paragraphs(doc) -> list[str]:
    """
    Extract text from all paragraphs in the document.
    Skips empty paragraphs.

    Args:
        doc: python-docx Document object

    Returns:
        list[str]: List of non-empty paragraph texts
    """
    paragraphs = []

    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            paragraphs.append(text)

    logger.info("Extracted %d paragraphs", len(paragraphs))
    return paragraphs


# Table extraction

def _extract_tables(doc) -> list[str]:
    """
    Extract text from all tables in the document.
    Each row is joined with a pipe separator for readability.
    Tables are common in resumes for skills and experience layouts.

    Args:
        doc: python-docx Document object

    Returns:
        list[str]: List of table row texts
    """
    rows = []

    for table_idx, table in enumerate(doc.tables):
        for row in table.rows:
            cell_texts = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cell_texts:
                rows.append(" | ".join(cell_texts))

    logger.info("Extracted %d table rows from %d tables", len(rows), len(doc.tables))
    return rows


# Header and footer extraction

def _extract_headers_footers(doc) -> list[str]:
    """
    Extract text from headers and footers across all sections.
    Headers often contain candidate name and contact info in resumes.

    Args:
        doc: python-docx Document object

    Returns:
        list[str]: List of non-empty header and footer texts
    """
    texts = []

    for section in doc.sections:
        for hf in [section.header, section.footer]:
            if hf:
                for para in hf.paragraphs:
                    text = para.text.strip()
                    if text:
                        texts.append(text)

    logger.info("Extracted %d header/footer lines", len(texts))
    return texts


# Public API

def parse_docx(file_path: str) -> str:
    """
    Extract text from a DOCX file using python-docx.
    Extracts paragraphs, tables, and headers/footers.

    Args:
        file_path (str): Path to DOCX or DOC file

    Returns:
        str: Extracted text content

    Raises:
        FileNotFoundError : File does not exist
        ValueError        : Unsupported file format
        RuntimeError      : Extraction failed
    """
    path = Path(file_path)

    if not path.exists():
        raise FileNotFoundError(f"File not found: {file_path}")

    if path.suffix.lower() not in (".docx", ".doc"):
        raise ValueError(f"Expected a DOCX/DOC file, got: {path.suffix}")

    logger.info("Parsing DOCX: %s", path.name)

    try:
        from docx import Document

        doc = Document(str(path))

        paragraphs     = _extract_paragraphs(doc)
        tables         = _extract_tables(doc)
        headers_footers = _extract_headers_footers(doc)

        # Combine all sections — headers first as they often have contact info
        all_text = "\n".join(headers_footers + paragraphs + tables).strip()

        if not all_text:
            raise ValueError(
                f"No text could be extracted from: {path.name}. "
                "The file may be empty or use an unsupported format."
            )

        logger.info("DOCX extraction complete — %d chars extracted", len(all_text))
        return all_text

    except ImportError:
        raise RuntimeError(
            "python-docx is not installed. "
            "Run: pip install python-docx"
        )
    except ValueError:
        raise
    except Exception as e:
        logger.error("DOCX extraction failed: %s", e)
        raise RuntimeError(f"DOCX extraction failed: {e}") from e
