# api/routes/ats_maker.py
"""
/api/ats/* endpoints - ATS check, optimise, and DOCX export.
"""

import io

from docx import Document
from docx.shared import Pt
from fastapi import APIRouter, Depends, HTTPException
from fastapi.concurrency import run_in_threadpool
from fastapi.responses import JSONResponse, StreamingResponse

from api.routes.auth import get_current_user, get_current_user_llm_limited
from core import uploads
from core.config import JD_MAX_CHARS
from core.logger import get_logger
from repositories import resume_repo as resume_store
from schemas.ats import AtsCheckRequest, AtsDocxRequest, AtsOptimiseRequest
from services.ats import ats_check, generate_ats_resume
from services.extractors.jd_extractor import extract_jd
from services.parsers import extract_all_text

logger = get_logger(__name__)

router = APIRouter()

# Body text size for generated DOCX files - standard resume body size.
_DOCX_BODY_PT = 11


def _resolve_resume_text(user_id: str, resume_id: str, tmp: str) -> str:
    """Return extracted text for a stored resume or a temp upload token.

    Raises HTTPException when neither source resolves to readable text.
    """
    path = None
    if resume_id:
        record = resume_store.get(user_id, resume_id)
        if not record:
            raise HTTPException(status_code=404, detail="resume_not_found")
        path = record["file_path"]
    elif tmp:
        path = uploads.resolve(user_id, tmp)
        if not path:
            raise HTTPException(status_code=404, detail="resume_not_found")
    else:
        raise HTTPException(status_code=400, detail="resume_required")

    text = extract_all_text(path)
    if not text:
        raise HTTPException(status_code=422, detail="could_not_read_resume")
    return text


@router.post("/check")
async def api_ats_check(
    body: AtsCheckRequest,
    current_user: dict = Depends(get_current_user_llm_limited),
) -> JSONResponse:
    """
    Lightweight ATS scan - no LLM, no keyword injection.

    Accepts {resume_id | tmp, jd?}. Returns section heading flags, formatting
    flags, keyword coverage (if JD provided), and a composite ATS score.
    """
    resume_text = await run_in_threadpool(
        _resolve_resume_text,
        current_user["id"],
        (body.resume_id or "").strip(),
        (body.tmp or "").strip(),
    )

    jd_text = (body.jd or "").strip()[:JD_MAX_CHARS]
    required_skills = None
    if jd_text and len(jd_text) >= 50:
        jd_json = await run_in_threadpool(extract_jd, jd_text)
        if jd_json:
            required_skills = jd_json.get("required_skills") or []

    result = await run_in_threadpool(ats_check, resume_text, required_skills)
    return JSONResponse({"ok": True, **result})


@router.post("/optimise")
async def api_ats_optimise(
    body: AtsOptimiseRequest,
    current_user: dict = Depends(get_current_user_llm_limited),
) -> JSONResponse:
    """
    Full ATS optimization pipeline with LLM.

    Accepts {resume_id | tmp, jd}. Returns a complete ATS-optimised resume
    with coverage before/after, section flags, and formatting warnings.
    """
    jd_text = (body.jd or "").strip()[:JD_MAX_CHARS]
    if len(jd_text) < 50:
        raise HTTPException(status_code=400, detail="jd_required")

    resume_text = await run_in_threadpool(
        _resolve_resume_text,
        current_user["id"],
        (body.resume_id or "").strip(),
        (body.tmp or "").strip(),
    )

    result = await run_in_threadpool(generate_ats_resume, resume_text, jd_text)
    if not result:
        raise HTTPException(status_code=503, detail="llm_unavailable")

    return JSONResponse({"ok": True, **result})


def _build_docx(parsed: dict) -> bytes:
    """Render the optimised-resume JSON into a simple ATS-friendly DOCX."""
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(_DOCX_BODY_PT)

    summary = (parsed.get("summary") or "").strip()
    if summary:
        doc.add_heading("Summary", level=1)
        doc.add_paragraph(summary)

    experience = parsed.get("experience") or []
    if experience:
        doc.add_heading("Work Experience", level=1)
        for job in experience:
            header = " | ".join(
                p for p in [job.get("title"), job.get("company"), job.get("dates")] if p
            )
            p = doc.add_paragraph()
            p.add_run(header).bold = True
            for b in job.get("bullets") or []:
                doc.add_paragraph(b, style="List Bullet")

    skills = parsed.get("skills") or []
    if skills:
        doc.add_heading("Skills", level=1)
        doc.add_paragraph(", ".join(skills))

    education = parsed.get("education") or []
    if education:
        doc.add_heading("Education", level=1)
        for edu in education:
            line = " - ".join(
                p for p in [edu.get("degree"), edu.get("institution")] if p
            )
            year = edu.get("year")
            doc.add_paragraph(f"{line} ({year})" if year else line)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


@router.post("/docx")
async def api_ats_docx(
    body: AtsDocxRequest,
    current_user: dict = Depends(get_current_user),
) -> StreamingResponse:
    """Download the optimised resume (JSON from /optimise) as a DOCX file."""
    if not body.resume:
        raise HTTPException(status_code=400, detail="resume_required")

    data = await run_in_threadpool(_build_docx, body.resume)
    return StreamingResponse(
        iter([data]),
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={
            "Content-Disposition": "attachment; filename=ats_optimised_resume.docx"
        },
    )
