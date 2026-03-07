from docx import Document


def parser_docx_direct(file_path):
    """
    Extract text content from a DOCX file using python-docx.
    Extracts from both paragraphs and tables to capture
    all resume content including structured skill tables.
    
    Args:
        file_path (str): Path to the DOCX file
        
    Returns:
        str: Extracted and cleaned text content
        
    Raises:
        FileNotFoundError : If file does not exist
        ValueError        : If document is empty
        RuntimeError      : If extraction fails
    """
    try:
        doc = Document(file_path)

        lines = []

        # Paragraphs
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                lines.append(text)

        # Tables
        for table in doc.tables:
            for row in table.rows:
                cells = []
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    if cell_text:
                        cells.append(cell_text)
                if cells:
                    lines.append(" | ".join(cells))

        # Headers and Footers
        for section in doc.sections:

            for para in section.header.paragraphs:
                text = para.text.strip()
                if text:
                    lines.append(text)

            for para in section.footer.paragraphs:
                text = para.text.strip()
                if text:
                    lines.append(text)

        # Remove duplicates while preserving order
        seen = set()
        clean_lines = []
        for line in lines:
            if line not in seen:
                seen.add(line)
                clean_lines.append(line)

        text = "\n".join(clean_lines)

        if len(text.strip()) < 100:
            raise ValueError(
                "DOCX appears to be empty or contains insufficient content. Upload proper file again"
            )

        return text.strip()

    except FileNotFoundError:
        raise FileNotFoundError(f"File not found: {file_path}")

    except ValueError as e:
        raise ValueError(str(e))

    except Exception as e:
        raise RuntimeError(f"Error extracting text from DOCX: {e}")
